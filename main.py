import html

import langcodes
import streamlit as st
from deep_translator import GoogleTranslator
from langdetect import DetectorFactory, LangDetectException, detect
from nltk.tokenize import TreebankWordDetokenizer, wordpunct_tokenize
from spellchecker import SpellChecker

from io import BytesIO
import docx
from pypdf import PdfReader
from transformers import pipeline

st.set_page_config(page_title="NLP Studio", page_icon="AI", layout="wide")

DetectorFactory.seed = 0
MIN_INPUT_LENGTH = 3

SPELL_LANGS = {
    "en", "es", "fr", "pt", "de",
    "ru", "ar", "eu", "lv", "nl"
}

TARGET_LANGS = {
    "Vietnamese": "vi",
    "English": "en",
    "French": "fr",
    "Japanese": "ja",
    "Chinese": "zh-CN",
    "Korean": "ko",
    "Spanish": "es",
    "German": "de",
}

EXAMPLES_T = [
    "Every morning, I drink a cup of coffee.",
    "Bonjour, comment allez-vous?",
    "Xin chao, hom nay troi dep qua.",
]
EXAMPLES_S = [
    "Yesturday, I recieveed a mesage from my freind.",
    "Definately a great oppurtunity.",
    "Je voudraiis allerr au marchee.",
]

# --- CACHE FUNCTIONS ---
@st.cache_resource(show_spinner=False)
def get_spellchecker(code):
    return SpellChecker(language=code)

@st.cache_resource(show_spinner="Loading Summarization Model (First time might take a while)...")
def get_summarizer():
    # Sử dụng model distilbart nhỏ gọn, chất lượng tốt cho việc tóm tắt tiếng Anh
    return pipeline("summarization", model="sshleifer/distilbart-cnn-6-6")

# --- HELPER FUNCTIONS ---
def language_name(code):
    try:
        return langcodes.Language.get(code).display_name()
    except Exception:
        return code or "Unknown"

def detect_language(raw):
    try:
        return detect(raw)
    except LangDetectException:
        return None

def fix_typos(text, code):
    spell = get_spellchecker(code)
    tokens = wordpunct_tokenize(text)
    fixed = []
    for token in tokens:
        if token.isalpha() and len(token) > 1:
            suggestion = spell.correction(token.lower()) or token
            suggestion = suggestion.title() if token.istitle() else suggestion
            suggestion = suggestion.upper() if token.isupper() else suggestion
            fixed.append(suggestion)
        else:
            fixed.append(token)

    return TreebankWordDetokenizer().detokenize(fixed), fixed != tokens

def run_translation(text, target_code):
    raw = text.strip()
    if len(raw) < MIN_INPUT_LENGTH:
        return {"ok": False, "error": f"Please enter at least {MIN_INPUT_LENGTH} characters."}

    source = detect_language(raw)
    if source is None:
        return {"ok": False, "error": "Unable to detect the language."}

    if source == target_code:
        return {
            "ok": True,
            "source": language_name(source),
            "target": language_name(target_code),
            "translated": raw,
            "note": "The text is already in the target language.",
        }

    try:
        translated = GoogleTranslator(source=source, target=target_code).translate(raw)
    except Exception as e:
        return {"ok": False, "error": f"Translation error: {e}"}

    return {"ok": True, "source": language_name(source), "target": language_name(target_code), "translated": translated}

def run_spellcheck(text):
    raw = text.strip()
    if len(raw) < MIN_INPUT_LENGTH:
        return {"ok": False, "error": f"Please enter at least {MIN_INPUT_LENGTH} characters."}

    code = detect_language(raw)
    if code is None:
        return {"ok": False, "error": "Unable to detect the language."}

    if code not in SPELL_LANGS:
        return {"ok": False, "error": f"pyspellchecker does not support {language_name(code)} ({code}) yet."}

    fixed, changed = fix_typos(raw, code)
    return {"ok": True, "language": language_name(code), "fixed": fixed, "changed": changed}

# --- NEW FUNCTIONS FOR FILE PROCESSING & SUMMARIZATION ---
def extract_text_from_file(uploaded_file):
    """Hàm đọc chữ từ file TXT, DOCX, PDF"""
    filename = uploaded_file.name
    try:
        if filename.endswith('.txt'):
            return uploaded_file.read().decode("utf-8")
        elif filename.endswith('.docx'):
            doc = docx.Document(uploaded_file)
            return "\n".join([p.text for p in doc.paragraphs])
        elif filename.endswith('.pdf'):
            reader = PdfReader(uploaded_file)
            text = ""
            for page in reader.pages:
                text += page.extract_text() or ""
            return text
    except Exception as e:
        st.error(f"Error reading file: {e}")
        return None

def convert_to_docx(text):
    """Hàm bọc văn bản thành file Word để cho người dùng tải về"""
    doc = docx.Document()
    doc.add_paragraph(text)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


def inject_theme():
    st.markdown(
        """
        <style>
        :root {
            --app-bg: #f6f7fb;
            --surface: #ffffff;
            --surface-soft: #f3f6fb;
            --surface-tinted: #f9fbff;
            --ink: #0f172a;
            --muted: #5f6f86;
            --line: #dbe3ee;
            --accent: #2f6fed;
            --accent-soft: #eef4ff;
            --good: #087f6f;
            --good-soft: #edf9f6;
            --warn: #a15c07;
            --warn-soft: #fff6e8;
            --danger: #ff4b4b;
            --danger-hover: #e83e3e;
            --upload-bg: #ffffff;
            --upload-border: #7c9bd6;
            --shadow: 0 14px 34px rgba(15, 23, 42, 0.07);
        }

        .stApp {
            background: var(--app-bg);
            color: var(--ink);
        }

        .block-container {
            max-width: 1180px;
            padding-top: 2rem;
            padding-bottom: 4rem;
        }

        .hero {
            background: var(--surface);
            border: 1px solid var(--line);
            border-radius: 8px;
            box-shadow: var(--shadow);
            padding: 1.5rem;
            margin-bottom: 1.25rem;
        }

        .eyebrow {
            color: var(--accent);
            font-size: 0.78rem;
            font-weight: 700;
            letter-spacing: 0;
            margin-bottom: 0.45rem;
            text-transform: uppercase;
        }

        .hero h1 {
            color: var(--ink);
            font-size: clamp(2rem, 4vw, 3.45rem);
            line-height: 1;
            letter-spacing: 0;
            margin: 0;
        }

        .hero p {
            color: var(--muted);
            font-size: 1.05rem;
            margin: 0.7rem 0 1.1rem;
            max-width: 680px;
        }

        .hero-grid {
            display: grid;
            grid-template-columns: repeat(3, minmax(0, 1fr));
            gap: 0.75rem;
        }

        .stat-card, .result-panel, .empty-panel {
            background: var(--surface);
            border: 1px solid var(--line);
            border-radius: 8px;
            padding: 1rem;
        }

        .stat-card strong {
            display: block;
            color: var(--ink);
            font-size: 1rem;
        }

        .stat-card span, .panel-meta {
            color: var(--muted);
            font-size: 0.88rem;
        }

        .result-panel {
            border-left: 4px solid var(--accent);
            box-shadow: 0 12px 30px rgba(15, 23, 42, 0.05);
            margin-top: 0.75rem;
        }

        .result-panel.success {
            background: var(--good-soft);
            border-left-color: var(--good);
        }

        .result-panel.warning {
            background: var(--warn-soft);
            border-left-color: var(--warn);
        }

        .result-panel h3, .empty-panel h3 {
            color: var(--ink);
            font-size: 1rem;
            margin: 0 0 0.45rem;
        }

        .result-panel p, .empty-panel p {
            color: var(--muted);
            margin: 0;
            white-space: pre-wrap;
        }

        .example-pill {
            background: var(--surface-tinted);
            border: 1px solid var(--line);
            border-radius: 8px;
            color: var(--ink);
            display: block;
            margin-bottom: 0.5rem;
            padding: 0.6rem 0.75rem;
        }

        div[data-testid="stTabs"] [role="tablist"] {
            background: #e9eef7;
            border: 1px solid var(--line);
            border-radius: 8px;
            display: inline-flex;
            gap: 0.55rem;
            padding: 0.25rem;
        }

        div[data-testid="stTabs"] button {
            background: transparent;
            border-radius: 8px;
            color: #334155 !important;
            font-weight: 650;
            padding-left: 0.9rem;
            padding-right: 0.9rem;
        }

        div[data-testid="stTabs"] button[aria-selected="true"] {
            background: var(--surface) !important;
            box-shadow: 0 2px 8px rgba(15, 23, 42, 0.08);
            color: var(--accent) !important;
        }

        div[data-testid="stTabs"] button[role="tab"] p {
            color: #334155 !important;
            font-weight: 700;
        }

        div[data-testid="stTabs"] button[aria-selected="true"] p {
            color: var(--accent) !important;
        }

        div[data-testid="stTextArea"] textarea,
        div[data-testid="stSelectbox"] div[data-baseweb="select"] {
            background: var(--surface) !important;
            border-radius: 8px;
            color: var(--ink) !important;
        }

        div[data-testid="stTextArea"] textarea {
            caret-color: var(--accent) !important;
        }

        div[data-testid="stTextArea"] textarea:focus {
            border-color: var(--accent) !important;
            box-shadow: 0 0 0 3px rgba(47, 111, 237, 0.16) !important;
            outline: none !important;
        }

        div[data-testid="stTextArea"] textarea::placeholder {
            color: #64748b !important;
            opacity: 1 !important;
        }

        div[data-testid="stTextArea"] label,
        div[data-testid="stTextArea"] label p,
        div[data-testid="stSelectbox"] label,
        div[data-testid="stSelectbox"] label p,
        div[data-testid="stFileUploader"] label,
        div[data-testid="stFileUploader"] label p {
            color: #334155 !important;
            font-weight: 650;
        }

        div[data-testid="stFileUploader"] {
            background: var(--surface) !important;
            border: 1px solid var(--upload-border);
            border-radius: 8px;
            padding: 0.75rem;
            box-shadow: 0 8px 22px rgba(47, 111, 237, 0.08);
        }

        div[data-testid="stFileUploaderDropzone"],
        section[data-testid="stFileUploaderDropzone"] {
            background: var(--upload-bg) !important;
            border: 1.5px dashed var(--upload-border) !important;
            border-radius: 8px !important;
            color: var(--ink) !important;
        }

        div[data-testid="stFileUploaderDropzone"] *,
        section[data-testid="stFileUploaderDropzone"] * {
            background-color: transparent !important;
            color: var(--ink) !important;
        }

        div[data-testid="stFileUploaderDropzone"] p,
        section[data-testid="stFileUploaderDropzone"] p,
        div[data-testid="stFileUploaderDropzone"] small,
        section[data-testid="stFileUploaderDropzone"] small,
        div[data-testid="stFileUploaderDropzone"] span {
            color: var(--ink) !important;
            opacity: 1 !important;
        }

        section[data-testid="stFileUploaderDropzone"] span {
            color: var(--ink) !important;
            opacity: 1 !important;
        }

        div[data-testid="stFileUploaderDropzone"] svg,
        section[data-testid="stFileUploaderDropzone"] svg {
            color: var(--accent) !important;
            fill: var(--accent) !important;
        }

        div[data-testid="stFileUploaderDropzone"] button,
        section[data-testid="stFileUploaderDropzone"] button {
            background: var(--accent) !important;
            border: 1px solid var(--accent) !important;
            color: #ffffff !important;
            font-weight: 700;
        }

        div[data-testid="stFileUploaderDropzone"] button *,
        section[data-testid="stFileUploaderDropzone"] button * {
            color: #ffffff !important;
        }

        div[data-testid="stExpander"] {
            background: var(--surface) !important;
            border: 1px solid var(--line);
            border-radius: 8px;
            box-shadow: none;
            overflow: hidden;
        }

        div[data-testid="stExpander"] details {
            background: var(--surface) !important;
        }

        div[data-testid="stExpander"] summary {
            background: var(--surface) !important;
            color: var(--ink) !important;
        }

        div[data-testid="stExpander"] summary:hover {
            background: var(--surface-tinted) !important;
        }

        div[data-testid="stExpander"] summary p,
        div[data-testid="stExpander"] summary svg {
            color: var(--ink) !important;
            fill: var(--ink) !important;
        }

        div[data-testid="stExpander"] div[data-testid="stText"],
        div[data-testid="stText"] {
            background: #ffffff !important;
            color: var(--ink) !important;
        }

        div[data-testid="stText"] pre,
        div[data-testid="stText"] code {
            background: #ffffff !important;
            color: #1e293b !important;
            opacity: 1 !important;
            text-shadow: none !important;
            white-space: pre-wrap !important;
        }

        div[data-testid="stExpander"] div,
        div[data-testid="stExpander"] p,
        div[data-testid="stExpander"] span {
            color: var(--ink) !important;
            opacity: 1 !important;
        }

        div[data-testid="stRadio"] label,
        div[data-testid="stRadio"] label p,
        div[data-testid="stRadio"] div[role="radiogroup"] label span {
            color: var(--ink) !important;
            opacity: 1 !important;
        }

        .stButton button, .stDownloadButton button {
            border-radius: 8px;
            color: var(--ink) !important;
            font-weight: 700;
            min-height: 2.85rem;
        }

        .stButton button[kind="primary"],
        .stDownloadButton button[kind="primary"] {
            background: var(--danger) !important;
            border: 1px solid var(--danger) !important;
            color: #ffffff !important;
        }

        .stButton button[kind="primary"]:hover,
        .stDownloadButton button[kind="primary"]:hover {
            background: var(--danger-hover) !important;
            border-color: var(--danger-hover) !important;
            color: #ffffff !important;
        }

        @media (max-width: 760px) {
            .hero-grid {
                grid-template-columns: 1fr;
            }
            .hero {
                padding: 1.1rem;
            }
        }
        </style>
        """,
        unsafe_allow_html=True,
    )


def render_hero():
    st.markdown(
        """
        <section class="hero">
            <div class="eyebrow">Language Intelligence Workspace</div>
            <h1>NLP Studio</h1>
            <p>Smart text tools for translation, correction, and document workflows.</p>
            <div class="hero-grid">
                <div class="stat-card"><strong>8 targets</strong><span>Translation languages ready</span></div>
                <div class="stat-card"><strong>10 spellcheckers</strong><span>Automatic language detection</span></div>
                <div class="stat-card"><strong>3 file types</strong><span>TXT, DOCX, and PDF support</span></div>
            </div>
        </section>
        """,
        unsafe_allow_html=True,
    )


def render_examples(examples):
    for example in examples:
        st.markdown(f'<span class="example-pill">{example}</span>', unsafe_allow_html=True)


def render_result_panel(title, body, meta=None, tone="success"):
    safe_title = html.escape(str(title))
    safe_body = html.escape(str(body))
    safe_meta = html.escape(str(meta)) if meta else ""
    meta_html = f'<div class="panel-meta">{safe_meta}</div>' if safe_meta else ""
    st.markdown(
        f"""
        <div class="result-panel {tone}">
            <h3>{safe_title}</h3>
            {meta_html}
            <p>{safe_body}</p>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_empty_panel(title, body):
    safe_title = html.escape(str(title))
    safe_body = html.escape(str(body))
    st.markdown(
        f"""
        <div class="empty-panel">
            <h3>{safe_title}</h3>
            <p>{safe_body}</p>
        </div>
        """,
        unsafe_allow_html=True,
    )


# --- STREAMLIT UI ---
inject_theme()
render_hero()
tab_t, tab_s, tab_d = st.tabs(["Translate", "Correction", "Document Lab"])

# --- TAB 1: TRANSLATE ---
with tab_t:
    st.session_state.setdefault("res_t", None)
    input_col, result_col = st.columns([1.05, 0.95], gap="large")

    with input_col:
        st.subheader("Translate text")
        with st.expander("Try an example", expanded=False):
            render_examples(EXAMPLES_T)

        with st.form("form_translate"):
            text_t = st.text_area(
                "Source text",
                height=180,
                placeholder="Paste a sentence or paragraph in any language...",
            )
            target = st.selectbox("Target language", list(TARGET_LANGS.keys()))
            submitted_t = st.form_submit_button("Translate", type="primary", use_container_width=True)

    if submitted_t:
        st.session_state.res_t = run_translation(text_t, TARGET_LANGS[target])

    with result_col:
        st.subheader("Result")
        res = st.session_state.res_t
        if res:
            if res["ok"]:
                render_result_panel(
                    "Translation",
                    res["translated"],
                    meta=f"{res['source']} to {res['target']}",
                    tone="success",
                )
                if res.get("note"):
                    st.info(res["note"])
            else:
                render_result_panel("Needs attention", res["error"], tone="warning")
        else:
            render_empty_panel("Ready when you are", "Your translated text will appear here.")

# --- TAB 2: CORRECTION ---
with tab_s:
    st.session_state.setdefault("res_s", None)
    input_col, result_col = st.columns([1.05, 0.95], gap="large")

    with input_col:
        st.subheader("Correct spelling")
        st.caption(f"Supported language codes: {', '.join(sorted(SPELL_LANGS))}")
        with st.expander("Try an example", expanded=False):
            render_examples(EXAMPLES_S)

        with st.form("form_spell"):
            text_s = st.text_area(
                "Text to check",
                height=180,
                placeholder="Paste text with possible spelling mistakes...",
            )
            submitted_s = st.form_submit_button("Check spelling", type="primary", use_container_width=True)

    if submitted_s:
        st.session_state.res_s = run_spellcheck(text_s)

    with result_col:
        st.subheader("Result")
        res = st.session_state.res_s
        if res:
            if res["ok"]:
                status = "Spelling errors corrected" if res["changed"] else "No spelling errors detected"
                render_result_panel("Corrected text", res["fixed"], meta=f"{res['language']} - {status}")
            else:
                render_result_panel("Needs attention", res["error"], tone="warning")
        else:
            render_empty_panel("Clean copy lands here", "Run a check to see corrected text and language details.")

# --- TAB 3: DOCUMENT LAB (NEW!) ---
with tab_d:
    st.subheader("Document processing")
    st.caption("Upload TXT, DOCX, or PDF files to summarize English text or translate the extracted content.")

    upload_col, action_col = st.columns([0.95, 1.05], gap="large")

    with upload_col:
        uploaded_file = st.file_uploader("Choose a document", type=["txt", "docx", "pdf"])
        render_empty_panel("Workflow", "1. Upload a file\n2. Preview extracted text\n3. Choose summarize or translate\n4. Download the result")

    if uploaded_file is not None:
        # 1. Đọc và trích xuất text từ file
        with st.spinner("Extracting text from file..."):
            file_text = extract_text_from_file(uploaded_file)

        if file_text:
            with action_col:
                render_result_panel(
                    "File loaded",
                    f"{uploaded_file.name}",
                    meta=f"{len(file_text):,} characters extracted",
                )

                # Xem trước nội dung file (giới hạn 500 ký tự cho đỡ rối mắt)
                with st.expander("Preview extracted content", expanded=True):
                    st.text(file_text[:1000] + ("..." if len(file_text) > 1000 else ""))

                # 2. Form lựa chọn tính năng xử lý
                action = st.radio(
                    "Action",
                    ["Summarize Text", "Translate Entire File"],
                    horizontal=True,
                )

                if action == "Translate Entire File":
                    doc_target = st.selectbox("Translate file to", list(TARGET_LANGS.keys()), key="doc_lang")

                process_btn = st.button("Process document", type="primary", use_container_width=True)

            if process_btn:
                output_result = ""

                # THÀNH PHẦN AI TÓM TẮT
                if action == "Summarize Text":
                    # Nhận diện ngôn ngữ, hiện tại model distilbart tối ưu nhất cho tiếng Anh (en)
                    lang = detect_language(file_text)
                    if lang != 'en':
                        st.warning("⚠️ Note: The AI Summarizer works best with English text. Results for other languages may vary.")

                    with st.spinner("AI is summarizing your document... Please wait."):
                        try:
                            summarizer = get_summarizer()
                            # Chia nhỏ văn bản nếu quá dài (model giới hạn khoảng 1024 tokens)
                            chunk_size = 2000
                            chunks = [file_text[i:i+chunk_size] for i in range(0, len(file_text), chunk_size)]

                            summaries = []
                            for chunk in chunks[:3]: # Giới hạn tối đa 3 chunks để tránh quá tải CPU máy cá nhân
                                res = summarizer(chunk, max_length=130, min_length=30, do_sample=False)
                                summaries.append(res[0]['summary_text'])

                            output_result = "\n\n".join(summaries)
                            render_result_panel("Summary generated", output_result)
                        except Exception as e:
                            st.error(f"Summarization error: {e}")

                # THÀNH PHẦN DỊCH FILE
                elif action == "Translate Entire File":
                    with st.spinner("Translating document..."):
                        # Chia nhỏ đoạn văn để dịch tránh bị giới hạn ký tự API của Google
                        chunk_size = 4000
                        chunks = [file_text[i:i+chunk_size] for i in range(0, len(file_text), chunk_size)]

                        translated_chunks = []
                        for chunk in chunks:
                            res = run_translation(chunk, TARGET_LANGS[doc_target])
                            if res["ok"]:
                                translated_chunks.append(res["translated"])
                            else:
                                st.error(res["error"])
                                break

                        output_result = "\n\n".join(translated_chunks)
                        render_result_panel(
                            f"Document translated to {doc_target}",
                            output_result[:2000] + ("..." if len(output_result) > 2000 else ""),
                        )

                # 3. NÚT DOWNLOAD KẾT QUẢ
                if output_result:
                    st.subheader("Download processed document")

                    # Cho phép tải dạng .txt
                    st.download_button(
                        label="Download TXT",
                        data=output_result,
                        file_name=f"processed_{uploaded_file.name.split('.')[0]}.txt",
                        mime="text/plain",
                        use_container_width=True,
                    )

                    # Cho phép tải dạng .docx (Word)
                    docx_data = convert_to_docx(output_result)
                    st.download_button(
                        label="Download DOCX",
                        data=docx_data,
                        file_name=f"processed_{uploaded_file.name.split('.')[0]}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                    )
